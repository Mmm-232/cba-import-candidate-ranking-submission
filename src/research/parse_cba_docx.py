"""
Module: parse_cba_docx.py
Purpose: Parse cba.docx and extract CBA import player lists by season as structured raw and cleaned tables.
"""
from __future__ import annotations

import argparse
import logging
import re
from pathlib import Path

import pandas as pd
from docx import Document

try:
    from .utils import PROCESSED_DIR, RAW_DIR, configure_logging, ensure_data_dirs, normalise_player_name, player_name_key
except ImportError:  # Allows: python src/parse_cba_docx.py
    from utils import PROCESSED_DIR, RAW_DIR, configure_logging, ensure_data_dirs, normalise_player_name, player_name_key


LOGGER = logging.getLogger(__name__)
SEASON_PATTERN = re.compile(r"(20\d{2})\s*[–-]\s*(20\d{2})")


# 功能：从 Word 文档内容中识别各个 CBA 赛季标题。
def extract_seasons(document: Document) -> list[str]:
    seasons: list[str] = []
    for paragraph in document.paragraphs:
        match = SEASON_PATTERN.search(paragraph.text)
        if match:
            seasons.append(f"{match.group(1)}-{match.group(2)}")
    return seasons


# 功能：按赛季读取 Word 文档中的外援姓名并完成标准化。
def parse_cba_docx(docx_path: Path) -> pd.DataFrame:
    if not docx_path.exists():
        raise FileNotFoundError(f"Cannot find input file: {docx_path}")

    document = Document(docx_path)
    seasons = extract_seasons(document)
    if len(seasons) != len(document.tables):
        raise ValueError(
            f"Expected one table per season, found {len(seasons)} season headings and {len(document.tables)} tables."
        )

    records: list[dict[str, object]] = []
    for table_index, (season, table) in enumerate(zip(seasons, document.tables), start=1):
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) < 2 or not cells[1]:
                continue

            row_number = cells[0]
            raw_name = cells[1]
            clean_name = normalise_player_name(raw_name)
            records.append(
                {
                    "cba_season": season,
                    "player_name_raw": raw_name,
                    "player_name_clean": clean_name,
                    "player_name_key": player_name_key(clean_name),
                    "target_cba_join_season": season,
                    "source_row_number": row_number,
                    "doc_table_index": table_index,
                }
            )

    df = pd.DataFrame.from_records(records)
    if df.empty:
        raise ValueError("No player rows were extracted from the document.")

    LOGGER.info("Extracted %s player-season rows across %s seasons.", len(df), df["cba_season"].nunique())
    return df


# 功能：分别保存原始姓名表和清洗后的 CBA 标签表。
def write_outputs(df: pd.DataFrame, raw_path: Path, clean_path: Path) -> None:
    raw_path.parent.mkdir(parents=True, exist_ok=True)
    clean_path.parent.mkdir(parents=True, exist_ok=True)

    raw_cols = ["cba_season", "source_row_number", "player_name_raw", "doc_table_index"]
    clean_cols = [
        "cba_season",
        "player_name_raw",
        "player_name_clean",
        "player_name_key",
        "target_cba_join_season",
        "source_row_number",
        "doc_table_index",
    ]
    df[raw_cols].to_csv(raw_path, index=False)
    df[clean_cols].to_csv(clean_path, index=False)
    LOGGER.info("Wrote raw CBA imports to %s", raw_path)
    LOGGER.info("Wrote clean CBA imports to %s", clean_path)


# 功能：执行 Word 外援名单解析流程并保存结果。
def main() -> None:
    parser = argparse.ArgumentParser(description="Parse CBA foreign player lists from cba.docx.")
    parser.add_argument("--docx", type=Path, default=Path("cba.docx"), help="Path to the source Word document.")
    parser.add_argument("--raw-out", type=Path, default=RAW_DIR / "cba_imports_raw.csv")
    parser.add_argument("--clean-out", type=Path, default=PROCESSED_DIR / "cba_imports_clean.csv")
    parser.add_argument("--log-level", default="INFO")
    args = parser.parse_args()

    configure_logging(args.log_level)
    ensure_data_dirs()
    df = parse_cba_docx(args.docx)
    write_outputs(df, args.raw_out, args.clean_out)


if __name__ == "__main__":
    main()
